"""Render a workspace (ordered list of cards / analyticals) into a .docx.

Inverse of ``parser/extract.py``: that module reads runs and produces offset
spans; this one takes offset spans and produces runs. Highlights and
underlines round-trip; colors do not (the DB stores ``kind`` only, so all
highlights export as yellow and all underlines as ``single``).

The library takes plain DTOs rather than SQLAlchemy ORM objects so the
round-trip test can exercise it without a database.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


@dataclass
class ExportCard:
    tag: str
    tag_markup: list[dict]
    card_text: str
    markup: list[dict]
    cite_short: str
    raw_cite: str


@dataclass
class ExportAnalytical:
    argument: str
    argument_markup: list[dict]
    answer_to: str | None = None


@dataclass
class ExportEntry:
    header_path: list[str] = field(default_factory=list)
    card: ExportCard | None = None
    analytical: ExportAnalytical | None = None

    def __post_init__(self) -> None:
        if (self.card is None) == (self.analytical is None):
            raise ValueError("ExportEntry must have exactly one of card / analytical")


def render_workspace_to_docx(
    workspace_name: str, entries: list[ExportEntry]
) -> bytes:
    doc = Document()
    doc.core_properties.title = workspace_name

    prev_path: list[str] = []
    for entry in entries:
        path = entry.header_path or []
        common = 0
        while (
            common < len(prev_path)
            and common < len(path)
            and prev_path[common] == path[common]
        ):
            common += 1
        for level, header in enumerate(path[common:], start=common + 1):
            doc.add_heading(header, level=min(level, 9))
        prev_path = path

        if entry.card is not None:
            _emit_card(doc, entry.card)
        else:
            assert entry.analytical is not None
            _emit_analytical(doc, entry.analytical)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _emit_card(doc, card: ExportCard) -> None:
    tag_p = doc.add_paragraph()
    _add_runs(tag_p, card.tag, card.tag_markup, base_bold=True)

    cite_p = doc.add_paragraph()
    cite_run = cite_p.add_run(card.cite_short)
    cite_run.bold = True

    if card.raw_cite:
        doc.add_paragraph(card.raw_cite)

    body_p = doc.add_paragraph()
    _add_runs(body_p, card.card_text, card.markup)


def _emit_analytical(doc, analytical: ExportAnalytical) -> None:
    if analytical.answer_to:
        at_p = doc.add_paragraph()
        at_run = at_p.add_run(f"AT: {analytical.answer_to}")
        at_run.bold = True

    body_p = doc.add_paragraph()
    _add_runs(body_p, analytical.argument, analytical.argument_markup)


def _add_runs(
    paragraph,
    text: str,
    markup: list[dict],
    *,
    base_bold: bool = False,
) -> None:
    """Walk character-boundary segments and emit one run per segment.

    Mirrors the algorithm in web/render.py:_render_full so the export
    produces the same partition the HTML view shows. Highlight implies
    underline (per the highlight-subset-of-underline domain rule), so
    every highlighted run also gets ``run.underline = True`` to keep
    cards readable in plain Word without a Verbatim template.
    """
    if not text:
        return
    if not markup:
        run = paragraph.add_run(text)
        if base_bold:
            run.bold = True
        return

    boundaries = {0, len(text)}
    for span in markup:
        boundaries.add(span["start"])
        boundaries.add(span["end"])

    sorted_b = sorted(boundaries)
    for i in range(len(sorted_b) - 1):
        start, end = sorted_b[i], sorted_b[i + 1]
        if start >= end:
            continue
        underlined = any(
            s["kind"] == "underline" and s["start"] <= start and s["end"] >= end
            for s in markup
        )
        highlighted = any(
            s["kind"] == "highlight" and s["start"] <= start and s["end"] >= end
            for s in markup
        )
        run = paragraph.add_run(text[start:end])
        if base_bold:
            run.bold = True
        if underlined or highlighted:
            run.underline = True
        if highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
